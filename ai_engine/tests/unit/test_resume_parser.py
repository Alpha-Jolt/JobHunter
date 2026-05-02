"""Unit tests for resume parser."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest

from ai_engine.core.exceptions import FabricationDetectedError
from ai_engine.features.llm.base import LLMResult
from ai_engine.features.llm.prompting.prompt_loader import PromptLoader
from ai_engine.features.llm.prompting.registry import PromptRegistry
from ai_engine.features.resume.resume_parser import ResumeParser
from ai_engine.features.resume.strategies.llm_parser_strategy import LLMParserStrategy


def _make_parser(
    response_content: dict, prompts_dir: Path, allow_fabrication: bool = False
) -> ResumeParser:
    router = MagicMock()
    router.complete = AsyncMock(
        return_value=LLMResult(content=response_content, provider="mock", model="mock")
    )
    loader = PromptLoader(prompts_dir)
    registry = PromptRegistry(prompts_dir)
    strategy = LLMParserStrategy(router, loader, registry)
    return ResumeParser(strategy, allow_fabrication=allow_fabrication)


@pytest.mark.asyncio
async def test_resume_parser_parses_docx(tmp_path: Path, prompts_dir: Path):
    """Parser extracts and parses a DOCX resume."""
    # Create a minimal DOCX
    try:
        from docx import Document

        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("john@example.com")
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("Python, Docker")
        docx_path = tmp_path / "resume.docx"
        doc.save(str(docx_path))
    except ImportError:
        pytest.skip("python-docx not installed")

    response = {
        "personal": {"name": "John Doe", "email": "john@example.com"},
        "summary": "",
        "experience_entries": [],
        "education": [],
        "skills": ["Python", "Docker"],
        "certifications": [],
        "projects": [],
    }
    parser = _make_parser(response, prompts_dir)
    resume = await parser.parse(docx_path)
    assert "Python" in resume.skills


@pytest.mark.asyncio
async def test_resume_parser_detects_fabrication(tmp_path: Path, prompts_dir: Path):
    """Parser raises FabricationDetectedError when LLM invents skills."""
    try:
        from docx import Document

        doc = Document()
        doc.add_paragraph("Jane Smith")
        doc.add_paragraph("jane@example.com")
        doc.add_paragraph("SKILLS: Python")
        docx_path = tmp_path / "resume.docx"
        doc.save(str(docx_path))
    except ImportError:
        pytest.skip("python-docx not installed")

    response = {
        "personal": {"name": "Jane Smith", "email": "jane@example.com"},
        "summary": "",
        "experience_entries": [],
        "education": [],
        "skills": ["Python", "QuantumMachineLearningXYZ"],  # Fabricated
        "certifications": [],
        "projects": [],
    }
    parser = _make_parser(response, prompts_dir, allow_fabrication=False)

    with pytest.raises(FabricationDetectedError):
        await parser.parse(docx_path)
