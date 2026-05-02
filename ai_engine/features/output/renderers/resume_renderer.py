"""Resume renderer — generates DOCX and PDF from an OptimisedVariant."""

from __future__ import annotations

import subprocess
from pathlib import Path

from ai_engine.core.exceptions import OutputRenderError
from ai_engine.core.logging_.logger import get_logger
from ai_engine.features.optimization.models.optimised_variant import OptimisedVariant

logger = get_logger(__name__)


def _render_docx(variant: OptimisedVariant, output_path: Path) -> None:
    """Render variant to DOCX using python-docx.

    Args:
        variant: OptimisedVariant with resume content.
        output_path: Destination DOCX file path.
    """
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError as exc:
        raise OutputRenderError("python-docx not installed. Run: pip install python-docx") from exc

    doc = Document()

    # Header — name
    doc.add_heading("Resume", level=0)

    # Summary
    if variant.rewritten_summary:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(variant.rewritten_summary)

    # Skills
    if variant.prioritized_skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(variant.prioritized_skills))

    # Experience
    if variant.reordered_experience:
        doc.add_heading("Experience", level=1)
        for entry in variant.reordered_experience:
            doc.add_heading(f"{entry.role} — {entry.company}", level=2)
            if entry.duration:
                doc.add_paragraph(entry.duration)
            for resp in entry.responsibilities:
                doc.add_paragraph(resp, style="List Bullet")
            for ach in entry.achievements:
                doc.add_paragraph(ach, style="List Bullet")

    # Projects
    if variant.selected_projects:
        doc.add_heading("Projects", level=1)
        for proj in variant.selected_projects:
            doc.add_heading(proj.name, level=2)
            doc.add_paragraph(proj.description)
            if proj.technologies:
                doc.add_paragraph("Technologies: " + ", ".join(proj.technologies))

    # Certifications
    if variant.selected_certifications:
        doc.add_heading("Certifications", level=1)
        for cert in variant.selected_certifications:
            doc.add_paragraph(cert, style="List Bullet")

    doc.save(str(output_path))


def _render_pdf_from_docx(docx_path: Path, pdf_path: Path) -> None:
    """Convert DOCX to PDF using LibreOffice headless.

    Args:
        docx_path: Source DOCX file.
        pdf_path: Destination PDF file.

    Raises:
        OutputRenderError: If LibreOffice conversion fails.
    """
    try:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_path.parent),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
    except FileNotFoundError as exc:
        raise OutputRenderError(
            "LibreOffice not found. Install LibreOffice to enable PDF output."
        ) from exc
    if result.returncode != 0:
        raise OutputRenderError(f"LibreOffice PDF conversion failed: {result.stderr}")


async def render_resume(
    variant: OptimisedVariant,
    output_dir: Path,
) -> tuple[str, str]:
    """Render resume to DOCX and PDF.

    Args:
        variant: OptimisedVariant with resume content.
        output_dir: Directory to write output files.

    Returns:
        Tuple of (docx_path, pdf_path) as strings.

    Raises:
        OutputRenderError: If rendering fails.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"resume_{variant.job_id}.docx"
    pdf_path = output_dir / f"resume_{variant.job_id}.pdf"

    try:
        _render_docx(variant, docx_path)
        logger.info("resume_renderer.docx_written", path=str(docx_path))
    except Exception as exc:
        raise OutputRenderError(f"DOCX rendering failed: {exc}") from exc

    try:
        _render_pdf_from_docx(docx_path, pdf_path)
        logger.info("resume_renderer.pdf_written", path=str(pdf_path))
    except OutputRenderError:
        logger.warning("resume_renderer.pdf_skipped", reason="LibreOffice unavailable or failed")
        pdf_path = Path("")  # PDF is optional if LibreOffice is not installed

    return str(docx_path), str(pdf_path)
